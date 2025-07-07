import os
from io import BytesIO

import streamlit as st
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from openai import OpenAI


# (5) MS워드 문서 변환 함수 정의
def markdown_to_docx(markdown_content: str, font_name: str, base_font_size: int):
    doc = Document()
    lines = markdown_content.split("\n")
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if line.startswith("## "):
            paragraph = doc.add_paragraph()  # 새 문단(paragraph) 추가
            run = paragraph.add_run(line[3:])  # '## ' 이후 텍스트 추가
            font = run.font  # run 객체의 font 속성에 접근
            font.size = Pt(base_font_size + 3)  # 폰트 크기 설정
            font.name = font_name  # 폰트 종류 설정
            font.bold = True  # 폰트 굵기 설정
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)  # 한글 폰트 설정
        elif line.startswith("### "):
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(line[4:])
            font = run.font
            font.size = Pt(base_font_size + 1)
            font.name = font_name
            font.bold = True
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        else:
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(line)
            font = run.font
            font.size = Pt(base_font_size)
            font.name = font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    byte_io = BytesIO()  # 메모리상에서 바이트 데이터를 저장할 객체 생성
    doc.save(byte_io)  # MS워드 문서(docx)를 바이트 스트림(BytesIO)에 저장
    byte_io.seek(0)  # 스트림의 위치를 처음(0)으로 이동
    return byte_io  # 바이트 데이터 반환


def main():
    st.set_page_config(layout="wide")
    st.title("보고서 작성 프로그램")
    with st.sidebar:
        ai_test_api_key = os.environ.get("OPENAI_API_KEY")
        openai_api_key = st.text_input("OpenAI API Key", type="password", value=ai_test_api_key)
        # (6) 폰트 종류 및 크기 선택 위젯 추가
        font_name = st.selectbox("글꼴 선택:", ["맑은 고딕", "바탕체"])
        base_font_size = st.slider("기본 글자 크기 (pt):", value=11)
        # (1) OpenAI 클라이언트 생성
        if openai_api_key:
            client = OpenAI(api_key=openai_api_key)

        # (2) 보고서 작성 함수 정의
        def process_text(prompt, text):
            content = prompt + "\n" + text
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": content}],
            )
            return response.choices[0].message.content
    # (3) 보고서 작성을 위한 프롬프트 입력
    prompt = """
    너는 보고서 작성 전문가야.
    다음 형식으로 보고서를 작성해줘.
    - 마크다운을 활용해 체계적으로 작성할 것
    - heading2(##) 3개, 각 heading2 내에서는 heading3(###) 2개로 구성할 것
    - heading2의 내용은 300자 이상으로 작성할 것
    - 목차는 제외할 것
    - 보고서 내용만 응답 결과로 보여줄 것
    """
    default_user_input = """생성형 AI가 세상을 어떻게 바꿀 수 있을까?"""
    user_input = st.text_area(
        "작성할 보고서의 주제 또는 내용을 입력하세요:",
        value=default_user_input,
        height=70,
    )
    # (4) 보고서 작성
    if st.button("보고서 작성"):
        if not openai_api_key:
            st.info("계속하려면 OpenAI API Key를 추가하세요.")
            st.stop()
        if not user_input.strip():
            st.warning("작성할 보고서의 주제를 입력하세요.")
            st.stop()
        with st.spinner("작성 중..."):
            result = process_text(prompt, user_input)
            # (7) 필요 없는 문구 삭제
            # print(result)
            st.write(result)
            # (8) MS워드 문서 변환 함수 호출
            docx_file = markdown_to_docx(result, font_name, base_font_size)
            # (9) 다운로드 버튼 생성
            st.download_button(
                label="보고서 다운로드",
                data=docx_file,
                file_name="보고서.docx",
            )


if __name__ == "__main__":
    main()
