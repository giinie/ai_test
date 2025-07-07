from io import BytesIO

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

markdown_input = '''
## 생성형 AI의 개요 및 발전

### 생성형 AI의 정의
생성형 AI(Generative AI)는 주어진 데이터로부터 새로운 콘텐츠를 생성할 수 있는 알고리즘과 모델을 의미합니다. 이러한 AI는 텍스트, 이미지, 음악, 비디오 등 다양한 형태의 데이터를 생성할 수 있으며, 최신 기술의 발전과 더불어 그 가능성이 무궁무진하게 확대되고 있습니다. 특히, 대화형 AI와 딥러닝 기술의 발전은 생성형 AI의 성능을 크게 향상시켰고, 이는 산업 전반에 걸쳐 많은 혁신을 가져오는 계 기가 되고 있습니다.

### 생성형 AI의 기술적 발전
생성형 AI의 발전은 주로 인공지능의 기초가 되는 머신러닝과 딥러닝의 발전에 기인합니다. 특히, Transform 모델과 GAN(Generative Adversarial Network) 등이 대표적입니다. 이러한 기술들은 데이터를 이 해하고, 이를 바탕으로 새로운 콘텐츠를 생성하는 능력을 가지고 있어, 창작뿐만 아니라 미래 예측, 데이터 증강 등의 분야에서도 응용됩니다. 또한, 점점 더 많은 기업들이 이러한 기술을 상용화하여 경쟁 력을 강화하려는 노력을 기울이고 있습니다.
'''


# (1) MS워드 문서 변환 함수 정의
def markdown_to_docx(markdown_content: str, font_name: str, base_font_size: int):
    doc = Document()
    # (2) 샘플 텍스트를 줄 단위로 분할
    lines = markdown_content.split("\n")
    for line in lines:
        line = line.strip()
        if not line:
            continue
        # (3) 각 줄별로 다른 서식 지정
        # '## '(heading2 제목)로 시작되는 줄 감지
        if line.startswith("## "):
            paragraph = doc.add_paragraph()  # 새 문단(paragraph) 추가
            run = paragraph.add_run(line[3:])  # '## ' 이후 텍스트 추가
            font = run.font  # run 객체의 font 속성에 접근
            font.size = Pt(base_font_size + 5)  # 폰트 크기 설정
            font.name = font_name  # 폰트 종류 설정
            font.bold = True  # 폰트 굵기 설정
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)  # 한글 폰트 설정
        # '### '(heading3 제목)로 시작되는 줄 감지
        elif line.startswith("### "):
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(line[4:])
            font = run.font
            font.size = Pt(base_font_size + 3)
            font.name = font_name
            font.bold = True
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        # 일반 문단
        else:
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(line)
            font = run.font
            font.size = Pt(base_font_size)
            font.name = font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    # (4) 문서 객체를 바이트 데이타로 반환
    byte_io = BytesIO()  # 메모리상에서 바이트 데이터를 저장할 객체 생성
    doc.save(byte_io)  # MS워드 문서(docx)를 바이트 스트림(BytesIO)에 저장
    byte_io.seek(0)  # 스트림의 위치를 처음(0)으로 이동
    return byte_io  # 바이트 데이터 반환


# (5) MS워드 문서 변환 함수 호출
font_name = "맑은 고딕"  # 한글 폰트 설정
base_font_size = 12  # 기본 폰트 크기 설정
docx_file = markdown_to_docx(markdown_input, font_name, base_font_size)

# (6) 생성된 파일 저장
with open("샘플워드.docx", "wb") as f:
    f.write(docx_file.read())
